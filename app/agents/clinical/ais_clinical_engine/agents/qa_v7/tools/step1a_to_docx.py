"""qa_v7 Step 1a — run a sample and emit a Word-doc report.

Dev-only. Runs Step 1a against a sample of the 500-question corpus
(offset configurable so we can hit a different slice on each run)
and writes a .docx report the reviewer can read on any machine.

Usage:
    python -m app.agents.clinical.ais_clinical_engine.agents.qa_v7.tools.step1a_to_docx \\
        --input "/path/to/test_questions_A 500.docx" \\
        --offset 5 --sample-every 10 \\
        --output "/tmp/v7_step1a_report.docx"
"""
from __future__ import annotations

import argparse
import asyncio
import json
import os
import sys
import time
from collections import Counter
from typing import Any, Dict, List


def _load_questions(path: str) -> List[str]:
    from docx import Document
    doc = Document(path)
    raw = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    cleaned: List[str] = []
    for line in raw:
        s = line
        idx = 0
        while idx < len(s) and s[idx].isdigit():
            idx += 1
        if (idx > 0 and idx < len(s) and s[idx] == "."
                and idx + 1 < len(s) and s[idx + 1] == " "):
            s = s[idx + 2:].strip()
        cleaned.append(s)
    # Drop title line (no question mark, everything else has one)
    if len(cleaned) >= 2:
        q_count = sum(1 for q in cleaned if "?" in q)
        if q_count >= len(cleaned) * 0.5 and "?" not in cleaned[0]:
            cleaned = cleaned[1:]
    return cleaned


def _build_client():
    from anthropic import Anthropic
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY not set.", file=sys.stderr)
        sys.exit(2)
    return Anthropic(api_key=api_key)


async def _run_parse(questions: List[str]) -> List[Dict[str, Any]]:
    from app.agents.clinical.ais_clinical_engine.agents.qa_v7.query_parser import (
        QueryParserV7,
    )
    client = _build_client()
    parser = QueryParserV7(nlp_client=client)
    out: List[Dict[str, Any]] = []
    total_in = 0
    total_out = 0
    start = time.time()
    for i, q in enumerate(questions, 1):
        t0 = time.time()
        parsed, usage = await parser.parse(q)
        elapsed = time.time() - t0
        total_in += usage.get("input_tokens", 0)
        total_out += usage.get("output_tokens", 0)
        out.append({
            "idx": i,
            "question": q,
            "parsed": parsed.to_dict(),
            "usage": usage,
            "elapsed_sec": round(elapsed, 2),
        })
        print(
            f"[{i:3}/{len(questions)}] scope={parsed.scope} "
            f"conf={parsed.extraction_confidence:.2f} "
            f"anchors={len(parsed.anchor_terms)} "
            f"vars={len(parsed.scenario_variables)} "
            f"({elapsed:.1f}s)",
            file=sys.stderr,
        )
    total = time.time() - start
    print(
        f"\nDone. {total:.1f}s elapsed. "
        f"input_tokens={total_in:,} output_tokens={total_out:,}",
        file=sys.stderr,
    )
    # Stash totals on last record for the report generator
    out.append({
        "_meta": True,
        "total_input_tokens": total_in,
        "total_output_tokens": total_out,
        "total_elapsed_sec": round(total, 1),
    })
    return out


def _fmt_dict(d: Dict[str, Any]) -> str:
    if not d:
        return "{}"
    return json.dumps(d, ensure_ascii=False)


def _write_docx(
    results: List[Dict[str, Any]], output_path: str, meta: Dict[str, Any],
) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("qa_v7 Step 1a — Parser Output Report", level=0)

    # Header summary
    doc.add_paragraph(
        f"Questions: {meta['n']}  |  "
        f"Input tokens: {meta['total_input_tokens']:,}  |  "
        f"Output tokens: {meta['total_output_tokens']:,}  |  "
        f"Elapsed: {meta['total_elapsed_sec']:.1f}s  |  "
        f"Est. cost: ~${meta['est_cost']:.2f}"
    )
    doc.add_paragraph(
        f"Sample: every {meta['sample_every']}th question starting "
        f"at offset {meta['offset']} of the {meta['corpus_size']}-question "
        f"corpus ({meta['input']})."
    )

    # Summary table
    doc.add_heading("Aggregate summary", level=1)
    agg = meta["aggregate"]
    p = doc.add_paragraph()
    p.add_run(f"Scope: ").bold = True
    p.add_run(
        f"in_scope={agg['scope_in']}  "
        f"out_of_scope={agg['scope_out']}"
    )
    p2 = doc.add_paragraph()
    p2.add_run(f"Confidence buckets: ").bold = True
    buckets = agg["conf_buckets"]
    p2.add_run(
        f"1.0={buckets['1.0']}  "
        f"0.9–0.95={buckets['0.9–0.95']}  "
        f"0.7–0.9={buckets['0.7–0.9']}  "
        f"0.5–0.7={buckets['0.5–0.7']}  "
        f"<0.5={buckets['<0.5']}"
    )
    p3 = doc.add_paragraph()
    p3.add_run(f"Clarifications populated: ").bold = True
    p3.add_run(str(agg["clarified"]))
    p4 = doc.add_paragraph()
    p4.add_run(f"Questions with 0 anchors (in_scope): ").bold = True
    p4.add_run(str(agg["empty_anchors_in_scope"]))
    p5 = doc.add_paragraph()
    p5.add_run(f"Questions with ≥1 scenario variable: ").bold = True
    p5.add_run(str(agg["has_scvars"]))

    # Per-question detail
    doc.add_heading("Per-question detail", level=1)

    for i, r in enumerate(results, 1):
        q = r["question"]
        p = r["parsed"]

        # Question heading
        heading = doc.add_paragraph()
        run = heading.add_run(f"Q{i}. ")
        run.bold = True
        run.font.size = Pt(12)
        run2 = heading.add_run(q)
        run2.font.size = Pt(12)

        # scope + confidence line
        line = doc.add_paragraph()
        run = line.add_run("scope: ")
        run.bold = True
        scope = p["scope"]
        sr = line.add_run(scope)
        if scope == "out_of_scope":
            sr.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
        line.add_run("   ")
        run = line.add_run("confidence: ")
        run.bold = True
        conf = p["extraction_confidence"]
        cr = line.add_run(f"{conf}")
        if conf < 0.7:
            cr.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)

        if p.get("clarification"):
            cp = doc.add_paragraph()
            run = cp.add_run("clarification: ")
            run.bold = True
            cp.add_run(p["clarification"])

        # Summary
        sp = doc.add_paragraph()
        run = sp.add_run("question_summary: ")
        run.bold = True
        sp.add_run(p["question_summary"])

        # Anchors
        ap = doc.add_paragraph()
        run = ap.add_run("anchor_terms: ")
        run.bold = True
        ap.add_run(_fmt_dict(p.get("anchor_terms") or {}))

        # Scenario vars
        vp = doc.add_paragraph()
        run = vp.add_run("scenario_variables: ")
        run.bold = True
        vp.add_run(_fmt_dict(p.get("scenario_variables") or {}))

        # Intent placeholders (Step 1b not built)
        ip = doc.add_paragraph()
        run = ip.add_run("intent: ")
        run.bold = True
        ir = ip.add_run(str(p.get("intent")))
        if p.get("intent") is None:
            ir.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            ip.add_run(
                "   (Step 1b not built — populated by embedding "
                "classifier when that step lands)"
            )

        # Separator
        doc.add_paragraph("—" * 40)

    doc.save(output_path)


def _aggregate(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    n = len(results)
    scope_in = 0
    scope_out = 0
    buckets = {"1.0": 0, "0.9–0.95": 0, "0.7–0.9": 0, "0.5–0.7": 0, "<0.5": 0}
    clarified = 0
    empty_anchors_in_scope = 0
    has_scvars = 0
    for r in results:
        p = r["parsed"]
        if p["scope"] == "in_scope":
            scope_in += 1
        else:
            scope_out += 1
        c = float(p["extraction_confidence"])
        if c >= 0.95:
            buckets["1.0"] += 1
        elif c >= 0.9:
            buckets["0.9–0.95"] += 1
        elif c >= 0.7:
            buckets["0.7–0.9"] += 1
        elif c >= 0.5:
            buckets["0.5–0.7"] += 1
        else:
            buckets["<0.5"] += 1
        if p.get("clarification"):
            clarified += 1
        if p["scope"] == "in_scope" and not (p.get("anchor_terms") or {}):
            empty_anchors_in_scope += 1
        if p.get("scenario_variables") or {}:
            has_scvars += 1
    return {
        "scope_in": scope_in,
        "scope_out": scope_out,
        "conf_buckets": buckets,
        "clarified": clarified,
        "empty_anchors_in_scope": empty_anchors_in_scope,
        "has_scvars": has_scvars,
    }


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--offset", type=int, default=0)
    ap.add_argument("--sample-every", type=int, default=10)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()

    all_questions = _load_questions(args.input)
    sampled = all_questions[args.offset::args.sample_every]
    print(
        f"Corpus: {len(all_questions)}  Sample: {len(sampled)} "
        f"(offset={args.offset}, every {args.sample_every})",
        file=sys.stderr,
    )

    results_plus_meta = asyncio.run(_run_parse(sampled))
    # Split off the _meta trailer
    meta_rec = results_plus_meta[-1]
    results = results_plus_meta[:-1]

    aggregate = _aggregate(results)
    est_cost = (
        meta_rec["total_input_tokens"] / 1_000_000 * 3
        + meta_rec["total_output_tokens"] / 1_000_000 * 15
    )
    report_meta = {
        "n": len(results),
        "total_input_tokens": meta_rec["total_input_tokens"],
        "total_output_tokens": meta_rec["total_output_tokens"],
        "total_elapsed_sec": meta_rec["total_elapsed_sec"],
        "est_cost": est_cost,
        "sample_every": args.sample_every,
        "offset": args.offset,
        "corpus_size": len(all_questions),
        "input": os.path.basename(args.input),
        "aggregate": aggregate,
    }
    _write_docx(results, args.output, report_meta)

    # Also drop raw JSON alongside for later analysis
    raw_path = args.output.replace(".docx", ".json")
    with open(raw_path, "w") as f:
        json.dump(
            {"meta": report_meta, "results": results}, f, indent=2,
        )

    print(
        f"\nWrote:\n  {args.output}\n  {raw_path}",
        file=sys.stderr,
    )


if __name__ == "__main__":
    main()
